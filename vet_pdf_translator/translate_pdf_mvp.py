#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""MVP-скрипт для локального перевода ветеринарного PDF через Ollama."""

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple

import fitz  # PyMuPDF
import requests
from docx import Document

OLLAMA_URL = "http://localhost:11434"
DEFAULT_GLOSSARY = {
    "canine": "собачий / относящийся к собакам",
    "feline": "кошачий / относящийся к кошкам",
    "bovine": "крупного рогатого скота",
    "equine": "лошадиный / относящийся к лошадям",
    "clinical signs": "клинические признаки",
    "diagnosis": "диагностика",
    "treatment": "лечение",
    "prognosis": "прогноз",
    "anthelmintic": "антигельминтный препарат",
    "dehydration": "обезвоживание",
    "vomiting": "рвота",
    "diarrhea": "диарея",
    "leukopenia": "лейкопения",
}

# Ищем числа, дозировки и единицы измерения в одной регулярке.
NUM_UNIT_RE = re.compile(
    r"(?i)\\b\\d+(?:[.,]\\d+)?\\s*(?:mg\\/kg|mg|g|kg|mcg|µg|ml|mL|L|IU|U|%|°C|°F|bpm|mmHg|q\\d+h|h|hr|hrs|day|days|week|weeks|month|months|year|years)\\b"
)
# Ищем латинские биномиальные названия вида Escherichia coli.
LATIN_BINOMIAL_RE = re.compile(r"\\b([A-Z][a-z]+\\s+[a-z]{2,})\\b")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Локальный перевод текстового PDF через Ollama")
    parser.add_argument("--pdf", required=True, help="Путь к исходному PDF")
    parser.add_argument("--model", default="qwen3:8b", help="Модель Ollama")
    parser.add_argument("--max-block-chars", type=int, default=2500, help="Максимум символов в чанке")
    parser.add_argument("--num-ctx", type=int, default=4096, help="Контекст модели")
    parser.add_argument("--temperature", type=float, default=0.1, help="Температура генерации")
    parser.add_argument("--force-extract", action="store_true", help="Пересоздать blocks.jsonl")
    parser.add_argument("--force-translate", action="store_true", help="Пересоздать translated.jsonl")
    return parser.parse_args()


def ensure_dirs() -> Tuple[Path, Path, Path]:
    base = Path(__file__).resolve().parent
    work_dir = base / "work"
    out_dir = base / "output"
    input_dir = base / "input"
    work_dir.mkdir(parents=True, exist_ok=True)
    out_dir.mkdir(parents=True, exist_ok=True)
    input_dir.mkdir(parents=True, exist_ok=True)
    return input_dir, work_dir, out_dir


def ensure_glossary(glossary_path: Path) -> Dict[str, str]:
    if not glossary_path.exists():
        with glossary_path.open("w", encoding="utf-8") as f:
            json.dump(DEFAULT_GLOSSARY, f, ensure_ascii=False, indent=2)
        print(f"[INFO] Создан glossary.json: {glossary_path}")

    with glossary_path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, dict):
        raise ValueError("glossary.json должен содержать JSON-объект (словарь)")
    return data


def check_ollama_and_model(model: str) -> None:
    try:
        resp = requests.get(f"{OLLAMA_URL}/api/tags", timeout=10)
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise RuntimeError(
            "Ollama недоступна по http://localhost:11434. Проверьте, что Ollama запущена."
        ) from exc

    try:
        probe = requests.post(
            f"{OLLAMA_URL}/api/generate",
            json={"model": model, "prompt": "ping", "stream": False},
            timeout=60,
        )
        probe.raise_for_status()
    except requests.RequestException as exc:
        raise RuntimeError(
            f"Модель '{model}' не отвечает. Проверьте, что она скачана (ollama pull {model})."
        ) from exc


def normalize_text(text: str) -> str:
    # Убираем лишние пробелы по краям, но переносы внутри оставляем.
    return "\n".join(line.rstrip() for line in text.strip().splitlines()).strip()


def split_long_text(text: str, max_chars: int) -> List[str]:
    # Если текст уже маленький, просто возвращаем как есть.
    if len(text) <= max_chars:
        return [text]

    chunks: List[str] = []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    current = ""
    for para in paragraphs:
        candidate = (current + "\n\n" + para).strip() if current else para
        if len(candidate) <= max_chars:
            current = candidate
            continue

        if current:
            chunks.append(current)
            current = ""

        # Если абзац всё ещё слишком длинный — делим по предложениям.
        if len(para) > max_chars:
            sentences = re.split(r"(?<=[.!?])\s+", para)
            sent_buf = ""
            for sent in sentences:
                sent = sent.strip()
                if not sent:
                    continue
                candidate_sent = (sent_buf + " " + sent).strip() if sent_buf else sent
                if len(candidate_sent) <= max_chars:
                    sent_buf = candidate_sent
                else:
                    if sent_buf:
                        chunks.append(sent_buf)
                    # Если одно предложение слишком длинное — режем по символам.
                    if len(sent) > max_chars:
                        for i in range(0, len(sent), max_chars):
                            chunks.append(sent[i : i + max_chars])
                        sent_buf = ""
                    else:
                        sent_buf = sent
            if sent_buf:
                chunks.append(sent_buf)
        else:
            current = para

    if current:
        chunks.append(current)
    return chunks


def extract_blocks(pdf_path: Path, max_block_chars: int) -> List[Dict]:
    blocks_out: List[Dict] = []

    with fitz.open(pdf_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            raw_blocks = page.get_text("blocks")
            # Сортируем блоки как в задании: сверху вниз, затем слева направо.
            raw_blocks = sorted(raw_blocks, key=lambda b: (b[1], b[0]))

            block_counter = 0
            for raw in raw_blocks:
                text = normalize_text(raw[4] if len(raw) > 4 else "")
                if not text or len(text) < 3:
                    continue

                block_counter += 1
                chunks = split_long_text(text, max_block_chars)
                for chunk_index, chunk in enumerate(chunks, start=1):
                    blocks_out.append(
                        {
                            "page": page_index,
                            "block_id": f"p{page_index}_b{block_counter}_c{chunk_index}",
                            "source": chunk,
                            "translation": "",
                            "status": "new",
                        }
                    )

    return blocks_out


def save_jsonl(path: Path, rows: List[Dict]) -> None:
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def load_jsonl(path: Path) -> List[Dict]:
    rows: List[Dict] = []
    if not path.exists():
        return rows
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            rows.append(json.loads(line))
    return rows


def glossary_to_text(glossary: Dict[str, str]) -> str:
    return "\n".join(f"- {k}: {v}" for k, v in glossary.items())


def build_prompt(source_text: str, glossary_text: str) -> str:
    return f"""Use non-thinking mode.
You are a professional veterinary and medical translator.

Translate the English veterinary text into Russian.

Strict rules:
1. Return only the Russian translation.
2. Do not explain anything.
3. Do not add comments.
4. Preserve all numbers, dosages, units, intervals and percentages exactly.
5. Preserve Latin names of species, organisms, pathogens and drugs.
6. Preserve headings, lists, numbering and line breaks where possible.
7. Use accurate veterinary terminology.
8. If a term is ambiguous, add [проверить термин] after it.
9. Do not omit any sentence.
10. Do not add information that is not present in the original.

Veterinary glossary:
{glossary_text}

Text:
{source_text}
"""


def translate_one_block(
    source_text: str,
    model: str,
    glossary_text: str,
    num_ctx: int,
    temperature: float,
) -> str:
    prompt = build_prompt(source_text, glossary_text)
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": temperature,
            "num_ctx": num_ctx,
            "top_p": 0.9,
            "repeat_penalty": 1.05,
        },
    }

    response = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=300)
    response.raise_for_status()
    data = response.json()
    text = (data.get("response") or "").strip()
    if not text:
        raise RuntimeError("Пустой ответ модели")
    return text


def translate_blocks(
    blocks: List[Dict],
    translated_path: Path,
    model: str,
    glossary: Dict[str, str],
    num_ctx: int,
    temperature: float,
    force_translate: bool,
) -> List[Dict]:
    if force_translate and translated_path.exists():
        translated_path.unlink()

    translated_rows = load_jsonl(translated_path)
    translated_map = {row["block_id"]: row for row in translated_rows}

    glossary_text = glossary_to_text(glossary)
    result_rows: List[Dict] = []

    for block in blocks:
        existing = translated_map.get(block["block_id"])
        if existing and existing.get("status") == "done":
            result_rows.append(existing)
            continue

        current = dict(block)
        try:
            translated = translate_one_block(
                source_text=block["source"],
                model=model,
                glossary_text=glossary_text,
                num_ctx=num_ctx,
                temperature=temperature,
            )
            current["translation"] = translated
            current["status"] = "done"
            current.pop("error", None)
            print(f"[OK] Переведён {block['block_id']}")
        except Exception as exc:  # noqa: BLE001
            # Не падаем целиком, а отмечаем ошибку и идём дальше.
            current["translation"] = ""
            current["status"] = "error"
            current["error"] = str(exc)
            print(f"[ERR] Ошибка в {block['block_id']}: {exc}")

        result_rows.append(current)
        # Сохраняем прогресс после каждого блока.
        save_jsonl(translated_path, result_rows)

    return result_rows


def extract_num_units(text: str) -> List[str]:
    return [m.group(0).strip() for m in NUM_UNIT_RE.finditer(text or "")]


def extract_latin_names(text: str) -> List[str]:
    # Берём только пары слов с заглавной + строчной формой.
    # Это простая эвристика для MVP.
    return [m.group(1) for m in LATIN_BINOMIAL_RE.finditer(text or "")]


def quality_check(blocks: List[Dict]) -> List[Dict]:
    report: List[Dict] = []

    for b in blocks:
        source = b.get("source", "")
        translation = b.get("translation", "")
        issues: List[str] = []

        src_nums = extract_num_units(source)
        tr_nums = extract_num_units(translation)
        if src_nums != tr_nums:
            issues.append("Числа/дозировки/единицы отличаются от оригинала")

        src_latin = extract_latin_names(source)
        missing_latin = [name for name in src_latin if name not in translation]
        if missing_latin:
            issues.append("Потеряны латинские названия: " + ", ".join(missing_latin))

        if "[проверить термин]" in translation:
            issues.append("Есть пометка [проверить термин]")

        if b.get("status") == "error":
            issues.append("Ошибка перевода блока")

        if issues:
            report.append(
                {
                    "page": b.get("page"),
                    "block_id": b.get("block_id"),
                    "source": source,
                    "translation": translation,
                    "issues": issues,
                }
            )

    return report


def build_docx(translated_blocks: List[Dict], qa_report: List[Dict], output_docx: Path) -> None:
    doc = Document()
    doc.add_heading("Перевод ветеринарного PDF", level=1)
    doc.add_paragraph(
        "Оригинал слева, перевод справа. Блоки с возможными проблемами отмечены строкой QA."
    )

    qa_ids = {item["block_id"] for item in qa_report}
    current_page = None

    for b in translated_blocks:
        page = b.get("page")
        if page != current_page:
            current_page = page
            doc.add_heading(f"Страница {page}", level=2)

        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = b.get("source", "")
        row.cells[1].text = b.get("translation", "")

        if b.get("block_id") in qa_ids:
            p = doc.add_paragraph()
            run = p.add_run("QA: проверить числа, единицы, латинские названия или терминологию.")
            run.bold = True

    doc.save(output_docx)


def main() -> int:
    args = parse_args()
    input_dir, work_dir, out_dir = ensure_dirs()

    glossary_path = Path(__file__).resolve().parent / "glossary.json"
    blocks_path = work_dir / "blocks.jsonl"
    translated_path = work_dir / "translated.jsonl"
    qa_path = work_dir / "qa_report.json"
    output_docx = out_dir / "translated.docx"

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        print(f"[FATAL] PDF не найден: {pdf_path}")
        print(f"Подсказка: положите файл в {input_dir / 'source.pdf'}")
        return 1

    try:
        glossary = ensure_glossary(glossary_path)
    except Exception as exc:  # noqa: BLE001
        print(f"[FATAL] Ошибка чтения glossary.json: {exc}")
        return 1

    try:
        check_ollama_and_model(args.model)
    except Exception as exc:  # noqa: BLE001
        print(f"[FATAL] {exc}")
        return 1

    try:
        if args.force_extract and blocks_path.exists():
            blocks_path.unlink()

        if not blocks_path.exists():
            blocks = extract_blocks(pdf_path, max_block_chars=args.max_block_chars)
            save_jsonl(blocks_path, blocks)
            print(f"[INFO] Создан {blocks_path} ({len(blocks)} блоков)")
        else:
            blocks = load_jsonl(blocks_path)
            print(f"[INFO] Используем существующий {blocks_path} ({len(blocks)} блоков)")
    except Exception as exc:  # noqa: BLE001
        print(f"[FATAL] Ошибка извлечения блоков из PDF: {exc}")
        return 1

    if not blocks:
        print("[FATAL] После извлечения не найдено текстовых блоков.")
        return 1

    translated_blocks = translate_blocks(
        blocks=blocks,
        translated_path=translated_path,
        model=args.model,
        glossary=glossary,
        num_ctx=args.num_ctx,
        temperature=args.temperature,
        force_translate=args.force_translate,
    )

    qa_report = quality_check(translated_blocks)
    with qa_path.open("w", encoding="utf-8") as f:
        json.dump(qa_report, f, ensure_ascii=False, indent=2)
    print(f"[INFO] QA-отчёт сохранён: {qa_path}")

    try:
        build_docx(translated_blocks, qa_report, output_docx)
        print(f"[INFO] DOCX сохранён: {output_docx}")
    except Exception as exc:  # noqa: BLE001
        print(f"[FATAL] Ошибка сборки DOCX: {exc}")
        return 1

    done = sum(1 for b in translated_blocks if b.get("status") == "done")
    err = sum(1 for b in translated_blocks if b.get("status") == "error")
    print(f"[DONE] Готово. Успешно: {done}, ошибок: {err}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
